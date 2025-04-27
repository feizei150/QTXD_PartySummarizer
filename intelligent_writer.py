import os
from flask import Flask, request, render_template_string, session, redirect, url_for, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx import Document as DocWriter
from ollama import Client

app = Flask(__name__)
app.secret_key = 'your-secret-key'
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
client = Client(host="http://localhost:11434")
MODEL = "deepseek-r1:7b"

HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="zh">
<head>
    <meta charset="UTF-8">
    <title>智能党建写作助手 V3</title>
    <script src="https://cdn.tailwindcss.com"></script>
</head>
<body class="bg-gray-50 p-8">
    <div class="max-w-5xl mx-auto bg-white shadow-md rounded-lg p-6">
        <h1 class="text-2xl font-bold mb-4">📘 智能党建写作助手 V3</h1>
        <form method="post" enctype="multipart/form-data">
            <div class="mb-4">
                <label class="block font-semibold">🧑 你的身份介绍：</label>
                <textarea name="identity" rows="2" class="w-full p-2 border rounded" required>{{ identity or '' }}</textarea>
            </div>
            <div class="mb-4">
                <label class="block font-semibold">📄 上传原始材料（支持多个 .docx / .txt）：</label>
                <input type="file" name="files" multiple class="w-full border rounded" accept=".docx,.txt">
            </div>
            <div class="mb-4">
                <label class="block font-semibold">✍️ 写作需求（希望生成的内容、风格等）：</label>
                <textarea name="instruction" rows="3" class="w-full p-2 border rounded" required>{{ instruction or '' }}</textarea>
            </div>
            <div class="mb-4">
                <label class="block font-semibold">📑 上传样式模板（支持多个 .docx）：</label>
                <input type="file" name="templates" multiple class="w-full border rounded" accept=".docx">
            </div>
            <div class="mb-4">
                <label class="inline-flex items-center">
                    <input type="checkbox" name="clear_session" value="true" class="mr-2">
                    🚪 退出智能党建助手并清空内容
                </label>
            </div>
            <button type="submit" class="bg-indigo-600 text-white px-4 py-2 rounded hover:bg-indigo-700">生成总结</button>
        </form>

        {% if filenames %}
        <div class="mt-6">
            <h2 class="text-lg font-semibold">📁 上传的原始材料：</h2>
            <ul class="list-disc pl-6">
                {% for file in filenames %}
                <li>{{ file }}</li>
                {% endfor %}
            </ul>
        </div>
        {% endif %}

        {% if result %}
        <div class="mt-8">
            <h2 class="text-xl font-bold mb-2">📄 总结结果：</h2>
            <pre class="bg-gray-100 p-4 rounded whitespace-pre-wrap">{{ result }}</pre>
            <a href="/download" class="inline-block mt-2 text-blue-600 underline">📥 下载 Word 版总结</a>
        </div>
        {% endif %}
    </div>
</body>
</html>
'''

def extract_text(path):
    if path.endswith(".docx"):
        doc = Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    elif path.endswith(".txt"):
        with open(path, encoding='utf-8') as f:
            return f.read()
    return ""

def extract_template_structure(template_path):
    try:
        doc = Document(template_path)
        sections = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 0]
        return "\n".join([f"【参考结构】{line}" for line in sections[:8]])
    except:
        return ""

@app.route("/", methods=["GET", "POST"])
def index():
    result = None
    filenames = []

    identity = session.get("identity")
    instruction = session.get("instruction")

    if request.method == "POST":
        if request.form.get("clear_session") == "true":
            session.clear()
            return redirect(url_for("index"))

        identity = request.form["identity"]
        instruction = request.form["instruction"]
        session["identity"] = identity
        session["instruction"] = instruction

        files = request.files.getlist("files")
        templates = request.files.getlist("templates")

        full_text = ""
        for file in files:
            filename = secure_filename(file.filename)
            if filename:
                filenames.append(filename)
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                file.save(filepath)
                full_text += extract_text(filepath) + "\n"

        template_structure = ""
        for template_file in templates:
            if template_file and template_file.filename.endswith(".docx"):
                template_path = os.path.join(UPLOAD_FOLDER, secure_filename(template_file.filename))
                template_file.save(template_path)
                template_structure += extract_template_structure(template_path) + "\n"

        prompt = f"""
你是一名党建写作智能助手，请根据以下身份信息、写作要求、原始材料内容，以及参考模板格式，生成一份结构清晰、语言规范、可直接使用的总结材料。

【用户身份】:
{identity}

【写作要求】:
{instruction}

{template_structure}

【原始材料内容】:
{full_text}
"""
        response = client.chat(model=MODEL, messages=[{"role": "user", "content": prompt}])
        result = response["message"]["content"]

        # 写入 Word 输出
        word_path = os.path.join(OUTPUT_FOLDER, "智能总结结果.docx")
        docx_writer = DocWriter()
        docx_writer.add_paragraph(result)
        docx_writer.save(word_path)

    return render_template_string(HTML_TEMPLATE, result=result, filenames=filenames, identity=identity, instruction=instruction)

@app.route("/download")
def download():
    return send_file(os.path.join(OUTPUT_FOLDER, "智能总结结果.docx"), as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True, port=7860)
